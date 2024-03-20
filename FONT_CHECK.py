from docx import Document


# def check_font_size(document, target_font_size):
#     for i, paragraph in enumerate(document.paragraphs):
#         for run in paragraph.runs:
#             if paragraph.style.font.size and paragraph.style.font.size != target_font_size:
#                 print(paragraph.style.font.size.pt)
#                 print(f"Несоответствие размера шрифта в абзаце: '{i}'")
#
#
# def check_font_style(document, target_font_style):
#     for i, paragraph in enumerate(document.paragraphs):
#         for run in paragraph.runs:
#             if run.font.name != target_font_style:
#                 print(f"Несоответствие стиля шрифта в абзаце: '{i}'")
#
#
# # Пример использования:
# if __name__ == "__main__":
#     # Замените путь к вашему документу
#     doc_path = "test.docx"
#
#     # Загрузка документа
#     doc = Document(doc_path)
#
#     # Проверка соответствия размера шрифта
#     target_font_size = 14
#     print(f"Проверка соответствия размера шрифта {target_font_size}...")
#     check_font_size(doc, target_font_size)
#
#     # Проверка соответствия стиля шрифта
#     target_font_style = "Times New Roman"
#     print(f"Проверка соответствия стиля шрифта {target_font_style}...")
#     check_font_style(doc, target_font_style)


def check_font_properties(docx_file):
    doc = Document(docx_file)

    # Задаем ожидаемые значения размера шрифта и стиля шрифта
    expected_font_size = 14
    expected_font_style = "Times New Roman"

    for i, paragraph in enumerate(doc.paragraphs):
        # print(paragraph.style.font.size)
        for run in paragraph.runs:
            # Проверяем соответствие размера шрифта
            if run.font.size is not None:
                if run.font.size.pt != expected_font_size:
                    print(f"Несоответствие размера шрифта в абзаце: {i}. Размер = {run.font.size.pt}")
            else:
                print(f"Размера шрифта в абзаце: {i} ОК. Размер = {run.font.size} (по умолчанию)")

            # Проверяем соответствие стиля шрифта
            if run.font.name is not None:
                if run.font.name != expected_font_style:
                    print(f"Несоответствие стиля шрифта в абзаце: {i}. Стиль = {run.font.name}")
            else:
                print(f"Стиль шрифта в абзаце: {i} ОК. Стиль = {paragraph.style.font.name} (по умолчанию)")


# Пример использования
check_font_properties("test.docx")
