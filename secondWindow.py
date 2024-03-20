import json
from PyQt5 import QtCore, QtGui
import docx
from docx import Document
from PyQt5.QtCore import QSize, Qt
from PyQt5.QtWidgets import (QApplication, QToolBar, QFileDialog, QLabel, QMainWindow, QPlainTextEdit, QVBoxLayout,
                             QWidget, QMessageBox, QPushButton, QComboBox, QScrollArea)
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ml import Ui_MainWindow
from file_reader import FileReader
import cfg as c
from docx_cls import FileManger
from constants import READ_ONLY, TITLE, SETTER, GOST

# from checkTitle import checkTitles
# from checkIndent import checkIndents
# from checkSetter import checkSetters
# from checkLineSpace import checkLineSpaces


GOST_FILE = "GOSTs.json"


class SecondWindow(QMainWindow):
    def __init__(self, MainClass):
        super().__init__()

        #   Начальные значения
        # self.gost = None
        self.pathFile = ''
        # self.paragraph_indent = ''
        # self.alignment = ''
        self.plain_text = None
        self.MainClass = MainClass
        # self.showMaximized()  # Полноэкранный режим

        self.setWindowTitle("Проверка по ГОСТам")
        self.setGeometry(100, 100, 700, 700)
        self.setMinimumSize(QSize(650, 700))

        self.title = QLabel("Проверка файла по ГОСТам", self)
        self.title.setGeometry(QtCore.QRect(150, 20, 400, 30))
        font = QtGui.QFont()
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(75)
        self.title.setFont(font)
        # self.title.setAlignment(QtCore.Qt.AlignCenter)

        labelGost = QLabel("Выберите ГОСТ", self)
        labelGost.setGeometry(30, 60, 200, 30)
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        labelGost.setFont(font)

        self.gostPicked = QLabel(self)
        self.gostPicked.setGeometry(220, 60, 180, 30)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.gostPicked.setFont(font)
        # self.gostPicked.setAlignment(QtCore.Qt.AlignCenter)

        self.choiceGost = QComboBox(self)
        self.choiceGost.setGeometry(30, 90, 180, 31)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.choiceGost.setFont(font)
        self.gost_keys = list(FileReader.get_files().keys())
        # self.gost_data = FileReader.read_file()    # Загрузка данных из JSON файла с использованием функции
        # self.gost_keys = list(self.gost_data.keys())   # Получение списка ключей
        self.choiceGost.addItems(self.gost_keys)  # Добавление ключей в QComboBox

        # Вывод параметров гостов
        self.fontStyleLabel = QLabel("Шрифт:", self)
        self.fontStyleLabel.setGeometry(30, 130, 400, 30)
        font.setPointSize(10)
        font.setBold(False)
        self.fontStyleLabel.setFont(font)

        self.fontSizeLabel = QLabel("Размер шрифта:", self)
        self.fontSizeLabel.setGeometry(30, 160, 400, 30)
        font.setPointSize(10)
        font.setBold(False)
        self.fontSizeLabel.setFont(font)

        self.paragraphIndentLabel = QLabel("Абзацный отступ:", self)
        self.paragraphIndentLabel.setGeometry(30, 190, 400, 30)
        font.setPointSize(10)
        font.setBold(False)
        self.paragraphIndentLabel.setFont(font)

        self.intervalLabel = QLabel("Межстрочный интервал:", self)
        self.intervalLabel.setGeometry(30, 220, 400, 30)
        font.setPointSize(10)
        font.setBold(False)
        self.intervalLabel.setFont(font)

        self.alignmentLabel = QLabel("Выравнивание:", self)
        self.alignmentLabel.setGeometry(30, 250, 400, 30)
        font.setPointSize(10)
        font.setBold(False)
        self.alignmentLabel.setFont(font)

        # Кнопка выбора файла
        self.pickFileButton = QPushButton("Выбрать файл", self)
        self.pickFileButton.setGeometry(30, 300, 161, 41)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.pickFileButton.setFont(font)

        self.filePicked = QLabel("", self)
        self.filePicked.setGeometry(220, 305, 700, 31)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.filePicked.setFont(font)
        # self.filePicked.setAlignment(QtCore.Qt.AlignCenter)

        self.checkFile = QPushButton("Проверить файл", self)
        self.checkFile.setGeometry(30, 350, 280, 35)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.checkFile.setFont(font)

        self.downloadFile = QPushButton("Скачать проверенный файл", self)
        self.downloadFile.setGeometry(30, 630, 280, 35)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.downloadFile.setFont(font)

        self.answer = QScrollArea(self)
        self.answer.setGeometry(20, 400, 550, 200)
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.answer.setFont(font)
        self.answer.setWidgetResizable(True)
        self.answer.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignTop)

        self.plain_text = QPlainTextEdit()
        self.plain_text.setReadOnly(READ_ONLY)

        layout = QVBoxLayout(self)
        layout.addWidget(self.plain_text)

        w = QWidget()
        w.setLayout(layout)
        self.answer.setWidget(w)

        #   Обработка событий
        # self.choiceGost.activated.connect(self.choiceGostActive)
        self.choiceGost.activated.connect(self.get_params_from_ghost)
        self.gostPicked.setText(self.choiceGost.currentText())
        self.pickFileButton.clicked.connect(self.pickFileButton_Clicked)
        self.checkFile.clicked.connect(self.checkFile_Clicked)
        self.downloadFile.clicked.connect(self.save_ready_file)

    # def choiceGostActive(self, index):
    #     selected_gost = self.choiceGost.itemText(index)
    #     properties = self.gost_data[selected_gost]
    #
    #     self.gostPicked.setText(selected_gost)
    #
    #     font_style = properties.get("font-style", "")
    #     font_size = properties.get("font-size", "")
    #     paragraph_indent = properties.get("paragraph-indent", "")
    #     interval = properties.get("interval", "")
    #     alignment = properties.get("alignment", "")
    #
    #     self.fontStyleLabel.setText(f"Font Style: {font_style}")
    #     self.fontSizeLabel.setText(f"Font Size: {font_size}")
    #     self.paragraphIndentLabel.setText(f"Paragraph Indent: {paragraph_indent}")
    #     self.intervalLabel.setText(f"Interval: {interval}")
    #     self.alignmentLabel.setText(f"Alignment: {alignment}")

    def get_params_from_ghost(self, index):
        # if self.gost in FileReader.get_files().keys():
        self.selected_gost = self.choiceGost.itemText(index)
        self.gostPicked.setText(self.selected_gost)
        params = FileReader(self.selected_gost + '.json').read_file()
        self.alignment = c.setter_gost[params['alignment']]
        self.indent = params['paragraph-indent']
        self.interval = params['interval']
        self.fname = params['font-style']
        self.fsize = params['font-size']

        self.fontStyleLabel.setText(f"Font Style: {self.fname}")
        self.fontSizeLabel.setText(f"Font Size: {self.fsize}")
        self.paragraphIndentLabel.setText(f"Paragraph Indent: {self.indent}")
        self.intervalLabel.setText(f"Interval: {self.interval}")
        self.alignmentLabel.setText(f"Alignment: {self.alignment}")

    # def loadGostValues(self):
    #     with open('GOSTs.json', 'r', encoding='utf-8') as file:
    #         gost_values = json.load(file)
    #     return gost_values

    def pickFileButton_Clicked(self):
        filename, filetype = QFileDialog.getOpenFileName(self,
                                                         "Выбрать файл",
                                                         '.',
                                                         'Word files (*.docx)')
        if filename == '':
            self.filePicked.setText('Файл не выбран.')
            self.pathFile = ''
        else:
            self.pathFile = filename
            filename = filename.split('/')[-1]
            self.filePicked.setText(filename)

    # def checkFile_Clicked(self, index):
    #     if self.pathFile == '':
    #         try:
    #             self.notSelectFile = QMessageBox()
    #             self.notSelectFile.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
    #             self.notSelectFile.setText('Вы не выбрали файл!')
    #             self.notSelectFile.setWindowTitle('Ошибка!')
    #             self.notSelectFile.setIcon(QMessageBox.Warning)
    #             res = self.notSelectFile.exec()
    #         except Exception as e:
    #             pass
    #     else:
    #         # if self.ui.enterIndent.text() == '':
    #         #     self.ui.enterIndent.setText('0')
    #         selected_gost = self.choiceGost.itemText(index)
    #         properties = self.gost_data[selected_gost]
    #
    #         # self.gostPicked.setText(selected_gost)
    #
    #         font_style = properties.get("font-style", "")
    #         font_size = properties.get("font-size", "")
    #         paragraph_indent = properties.get("paragraph-indent", "")
    #         interval = properties.get("interval", "")
    #         alignment = properties.get("alignment", "")
    #         document = Document(self.pathFile)
    #         text = checkIndents(paragraph_indent, document)
    #         text += checkSetters(alignment, document)
    #         text += checkLineSpaces(interval, document)
    #         # text += checkTitles(self.currentTitle.text(), document)
    #         document.save(self.pathFile)
    #         self.plain_text.setPlainText(text)

    def checkFile_Clicked(self):
        print(self.pathFile)
        if self.pathFile == '':
            try:
                self.notSelectFile = QMessageBox()
                self.notSelectFile.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
                self.notSelectFile.setText('Вы не выбрали файл!')
                self.notSelectFile.setWindowTitle('Ошибка!')
                self.notSelectFile.setIcon(QMessageBox.Warning)
                res = self.notSelectFile.exec()
            except Exception as e:
                pass
        else:
            self.selected_gost = self.gostPicked.text()
            print(self.selected_gost)
            # print(type(self.selected_gost))
            self.fileName = self.pathFile.split('/')[-1]
            self.path = f'./{self.fileName}'
            # print(self.path)
            obj = FileManger(docx.Document(self.path), gost=self.selected_gost, doc_rej=False)
            print(1)
            errors = obj.is_correct_document()
            print(2)
            self.plain_text.clear()
            self.plain_text.setPlainText(errors)

    def save_ready_file(self):
        obj2 = FileManger(docx.Document(self.path), gost=self.selected_gost, doc_rej=True)
        obj2.is_correct_document()


# Убрать потом
if __name__ == '__main__':
    app = QApplication([])
    window = SecondWindow(None)
    # window = SecondWindow(None)
    window.show()
    app.exec_()

# Оставить
# if __name__ == '__main__':
#     import sys
#     app = QApplication(sys.argv)
#     second_window = SecondWindow(None)
#     second_window.show()
#     sys.exit(app.exec_())
