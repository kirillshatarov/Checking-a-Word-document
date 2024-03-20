from docx import Document
from docx.shared import RGBColor
from constants2 import TITLE, INTERVALS, SETTER, CORRECTSEQUENCE
from docx.oxml import OxmlElement


# def add_comment(paragraph, comment_text):
#     comment = paragraph.add_comment(comment_text)
#     # Дополнительная настройка комментария, если необходимо
#     comment.author = 'Author'
#     comment.initials = 'AI'
#     comment.text = 'Additional comment text'
def checkDocument(document, currentTitle, currentIndent, currentSetter,
                  currentLineSpace):
    # цвета
    redTitle = RGBColor(255, 0, 0)
    greenSetter = RGBColor(0, 255, 0)
    blueIndent = RGBColor(0, 0, 255)
    yellowLineSpace = RGBColor(196, 120, 20)
    pinkParagraphSpace = RGBColor(250, 105, 165)

    # для заголовков
    titleOnRussia = currentTitle
    currentTitle = TITLE[currentTitle]
    count = 1

    # для отступа первой строки
    currentIndent = currentIndent.replace(',', '.')

    currentSetterRussian = currentSetter
    currentSetter = SETTER[currentSetter]

    currentLineSpace = currentLineSpace.replace(',', '.')

    text = 'ЗАГОЛОВКИ.\n'

    for paragraph in document.paragraphs:

        # Заголовки
        if paragraph.style.name in TITLE.values():
            try:
                if count != CORRECTSEQUENCE[paragraph.text.capitalize()]:
                    arr = []
                    for key, val in CORRECTSEQUENCE.items():
                        if val == count:
                            arr.append(key)
                    if arr:
                        for word in paragraph.runs:
                            word.font.color.rgb = redTitle
                        text += f'Неверный порядок заголовка.\nОжидалось: "{arr[0]}"\t Получили: "{paragraph.text}".\n'
                if paragraph.style.name == currentTitle:
                    text += f'Верно оформлен заголовок: "{paragraph.text}".\n'
                else:
                    c = True
                    for word in paragraph.runs:
                        word.font.color.rgb = redTitle
                        if c:
                            word.add_comment('Не тот стиль заголовка')
                            c = False
                    text += f'Заголовок "{paragraph.text}" оформлен в неверном стиле "{[key for key, val in TITLE.items() if paragraph.style.name == val][0]}".\n' \
                            f'Этот заголовок ожидался в стиле "{titleOnRussia}".\n'
            except KeyError:
                for word in paragraph.runs:
                    word.font.color.rgb = redTitle
                    word.add_comment(
                        f"Неверно написан заголовок. Ожидалось: {[key for key, value in CORRECTSEQUENCE.items() if value == count][0]}")
                text += f'"{paragraph.text}" не является стандартом госта. ' \
                        f'Ожидалось: "{[key for key, value in CORRECTSEQUENCE.items() if value == count][0]}".\n'
            finally:
                text += '-----\n'
                count += 1

        # Проверка отступа
        # Проверяем, не является ли абзац заголовком
        if paragraph.style.name not in TITLE.values():
            # проверяем, является ли абзац - списоком, если да, то идем дальше по документу.
            if len(paragraph._element.xpath('./w:pPr/w:numPr')) > 0:
                continue
            if paragraph.paragraph_format.first_line_indent is not None:
                # Округление так как в first_line_indent.cm не ровное число, а ~1.2505972222222221
                if round(paragraph.paragraph_format.first_line_indent.cm, 2) != float(currentIndent):
                    c = True
                    for word in paragraph.runs:
                        word.font.color.rgb = blueIndent
                        if c:
                            word.add_comment(
                                f'Отступ составляет {round(paragraph.paragraph_format.first_line_indent.cm, 2)} см.'
                                f'Необходим: {currentIndent}')
                            c = False

        # проверка выравнивания
        if paragraph.style.name not in TITLE.values():
            if currentSetter != paragraph.alignment and paragraph.alignment is not None:
                c = True
                for word in paragraph.runs:
                    word.font.color.rgb = greenSetter
                    if c:
                        word.add_comment(f'Выравнивание не соответствует критерию. Выравнивание '
                                         f'должно быть: {currentSetterRussian}')
                        c = False

        # проверка межстрочного интервала
        if paragraph.paragraph_format.line_spacing != float(currentLineSpace) and \
                paragraph is not None:
            c = True
            for word in paragraph.runs:
                word.font.color.rgb = yellowLineSpace
                if c:
                    word.add_comment(f'Межстрочный интервал не соответствует ГОСТу.'
                                     f'Он должен составлять: {currentLineSpace}')
                    c = False

        # if paragraph.paragraph_format.line_spacing != float(currentLineSpace) and \
        #         paragraph is not None:
        #     comment_text = f'Межстрочный интервал не соответствует ГОСТу. Он должен составлять: {currentLineSpace}'
        #     c = True
        #     for word in paragraph.runs:
        #         word.font.color.rgb = yellowLineSpace
        #         if c:
        #             add_comment(word, comment_text)
        #             c = False


filename = 'ML.docx'
document = Document(filename)
checkDocument(document, 'Заголовок', '1,25', 'по ширине', '1,5')
document.save(filename.split('.')[0] + 'Исправлен.docx')
